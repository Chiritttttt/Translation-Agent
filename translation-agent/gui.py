import os
import sys
import logging

# PyInstaller 打包后，工作目录可能不是 exe 所在目录
# 需要在加载 dotenv 之前定位 .env 文件
if getattr(sys, 'frozen', False):
    # onefile 模式：exe 解压目录
    _BASE_DIR = os.path.dirname(sys.executable)
else:
    _BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 依次搜索：exe 同目录 → 当前目录 → 脚本目录
for _env_dir in [_BASE_DIR, os.getcwd(), os.path.dirname(os.path.abspath(__file__))]:
    _env_path = os.path.join(_env_dir, '.env')
    if os.path.exists(_env_path):
        os.environ.setdefault('DOTENV_PATH', _env_path)
        break

from dotenv import load_dotenv
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QComboBox, QLineEdit, QTextEdit, QPushButton, QFileDialog,
    QTabWidget, QMessageBox, QGroupBox, QProgressBar, QDialog,
    QRadioButton, QButtonGroup, QFrame, QGraphicsDropShadowEffect, QSizePolicy,
    QHeaderView, QTableWidget, QTableWidgetItem, QGridLayout,
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize
from PyQt6.QtGui import QFont, QColor, QPalette, QIcon, QPixmap
import openai
import requests
from bs4 import BeautifulSoup

# 日志配置 — 打包后 console=False 也能写文件排查问题
_log_dir = os.path.join(_BASE_DIR, 'logs')
os.makedirs(_log_dir, exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(_log_dir, 'translation_agent.log'), encoding='utf-8'),
    ]
)
logger = logging.getLogger('TranslationAgent')

_dotenv_path = os.environ.get('DOTENV_PATH', None)
if _dotenv_path:
    load_dotenv(_dotenv_path)
    logger.info(f"已加载 .env: {_dotenv_path}")
else:
    load_dotenv()

_api_key = os.getenv("OPENAI_API_KEY", "")
_base_url = os.getenv("OPENAI_BASE_URL", "") or None
_model_name = os.getenv("OPENAI_MODEL", "deepseek-chat")

if not _api_key:
    logger.warning("OPENAI_API_KEY 未配置，请在 .env 中设置或启动后通过界面配置")

client = openai.OpenAI(
    api_key=_api_key or "sk-placeholder",
    base_url=_base_url,
)
model = _model_name


# ═══════════════════════════════════════════════════════════
# Arco Design 颜色系统
# ═══════════════════════════════════════════════════════════
class ArcoColors:
    """Arco Design 色板"""

    # ── 主色 ──
    PRIMARY          = "#165DFF"
    PRIMARY_LIGHT    = "#E8F3FF"
    PRIMARY_LIGHT2   = "#F2F7FF"
    PRIMARY_HOVER    = "#4080FF"
    PRIMARY_ACTIVE   = "#0E42D2"

    # ── 功能色 ──
    SUCCESS          = "#00B42A"
    SUCCESS_LIGHT    = "#E8FFEA"
    WARNING          = "#FF7D00"
    WARNING_LIGHT    = "#FFF7E8"
    DANGER           = "#F53F3F"
    DANGER_LIGHT     = "#FFECE8"
    INFO             = "#165DFF"
    INFO_LIGHT       = "#E8F3FF"

    # ── 中性色 ──
    TEXT_PRIMARY      = "#1D2129"
    TEXT_REGULAR      = "#4E5969"
    TEXT_SECONDARY    = "#86909C"
    TEXT_DISABLED     = "#C9CDD4"
    TEXT_CAPTION      = "#C9CDD4"

    BORDER           = "#E5E6EB"
    BORDER_LIGHT     = "#F2F3F5"
    FILL1            = "#F7F8FA"
    FILL2            = "#F2F3F5"
    FILL3            = "#E5E6EB"
    FILL4            = "#C9CDD4"

    BG_PAGE          = "#F2F3F5"
    BG_CARD          = "#FFFFFF"
    BG_ELEVATION     = "#FFFFFF"

    # ── 阴影 ──
    SHADOW_SM        = "0 1px 2px 0 rgba(0,0,0,0.03), 0 1px 6px -1px rgba(0,0,0,0.02), 0 2px 4px 0 rgba(0,0,0,0.02)"
    SHADOW_MD        = "0 3px 6px -4px rgba(0,0,0,0.12), 0 6px 16px 0 rgba(0,0,0,0.08), 0 9px 28px 8px rgba(0,0,0,0.05)"
    SHADOW_LG        = "0 6px 30px 5px rgba(0,0,0,0.05), 0 16px 24px 2px rgba(0,0,0,0.04), 0 8px 10px -5px rgba(0,0,0,0.08)"

    # ── 圆角 ──
    RADIUS_SM = "4px"
    RADIUS_MD = "8px"
    RADIUS_LG = "12px"

    # ── 字体 ──
    FONT_FAMILY = "'Microsoft YaHei UI', 'Microsoft YaHei', 'PingFang SC', 'Noto Sans SC', 'Segoe UI', sans-serif"
    FONT_FAMILY_EN = "'Segoe UI', 'Microsoft YaHei UI', -apple-system, sans-serif"


# ═══════════════════════════════════════════════════════════
# Arco Design QSS 样式表
# ═══════════════════════════════════════════════════════════
def arco_global_stylesheet():
    C = ArcoColors
    return f"""
    /* ── 全局 ── */
    QWidget {{
        font-family: {C.FONT_FAMILY};
        font-size: 13px;
        color: {C.TEXT_PRIMARY};
        font-weight: 400;
        letter-spacing: 0.2px;
    }}

    /* ── 主窗口 ── */
    QMainWindow {{
        background-color: {C.BG_PAGE};
    }}

    /* ── 卡片容器 (QGroupBox) ── */
    QGroupBox {{
        background: {C.BG_CARD};
        border: 1px solid {C.BORDER_LIGHT};
        border-radius: {C.RADIUS_LG};
        padding: 20px 20px 16px 20px;
        margin-top: 0px;
        margin-bottom: 4px;
    }}
    QGroupBox::title {{
        subcontrol-origin: margin;
        subcontrol-position: top left;
        color: {C.TEXT_PRIMARY};
        font-size: 14px;
        font-weight: 600;
        padding: 0 0 12px 0;
        left: 20px;
    }}

    /* ── 输入框 ── */
    QLineEdit {{
        background: {C.BG_CARD};
        border: 1px solid {C.BORDER};
        border-radius: {C.RADIUS_MD};
        padding: 8px 14px;
        font-size: 14px;
        color: {C.TEXT_PRIMARY};
        selection-background-color: {C.PRIMARY_LIGHT};
        selection-color: {C.PRIMARY};
    }}
    QLineEdit:hover {{
        border-color: {C.PRIMARY_HOVER};
    }}
    QLineEdit:focus {{
        border-color: {C.PRIMARY};
        border-width: 2px;
        padding: 7px 13px;
    }}
    QLineEdit:disabled {{
        background: {C.FILL1};
        color: {C.TEXT_DISABLED};
        border-color: {C.BORDER_LIGHT};
    }}

    /* ── 文本编辑框 ── */
    QTextEdit {{
        background: {C.BG_CARD};
        border: 1px solid {C.BORDER};
        border-radius: {C.RADIUS_MD};
        padding: 10px 14px;
        font-size: 14px;
        line-height: 1.6;
        color: {C.TEXT_PRIMARY};
        selection-background-color: {C.PRIMARY_LIGHT};
        selection-color: {C.PRIMARY};
    }}
    QTextEdit:hover {{
        border-color: {C.PRIMARY_HOVER};
    }}
    QTextEdit:focus {{
        border-color: {C.PRIMARY};
        border-width: 2px;
        padding: 9px 13px;
    }}
    QTextEdit:read-only {{
        background: {C.FILL1};
        color: {C.TEXT_REGULAR};
        border-color: {C.BORDER_LIGHT};
    }}

    /* ── 下拉框 ── */
    QComboBox {{
        background: {C.BG_CARD};
        border: 1px solid {C.BORDER};
        border-radius: {C.RADIUS_MD};
        padding: 8px 36px 8px 14px;
        font-size: 14px;
        color: {C.TEXT_PRIMARY};
        min-height: 22px;
    }}
    QComboBox:hover {{
        border-color: {C.PRIMARY_HOVER};
    }}
    QComboBox:focus, QComboBox::drop-down:pressed {{
        border-color: {C.PRIMARY};
    }}
    QComboBox::drop-down {{
        subcontrol-origin: padding;
        subcontrol-position: right center;
        width: 28px;
        border: none;
        background: transparent;
    }}
    QComboBox::down-arrow {{
        width: 12px;
        height: 12px;
        image: none;
        border-left: 5px solid transparent;
        border-right: 5px solid transparent;
        border-top: 6px solid {C.TEXT_SECONDARY};
    }}
    QComboBox QAbstractItemView {{
        background: {C.BG_ELEVATION};
        color: {C.TEXT_PRIMARY};
        border: 1px solid {C.BORDER};
        border-radius: {C.RADIUS_MD};
        padding: 4px 0;
        selection-background-color: {C.PRIMARY_LIGHT};
        selection-color: {C.PRIMARY};
        outline: none;
    }}
    QComboBox QAbstractItemView::item {{
        height: 36px;
        padding: 0 14px;
        color: {C.TEXT_PRIMARY};
    }}
    QComboBox QAbstractItemView::item:hover {{
        background-color: {C.FILL1};
    }}
    QComboBox QAbstractItemView::item:selected {{
        background-color: {C.PRIMARY_LIGHT};
        color: {C.PRIMARY};
    }}

    /* ── 标签页 ── */
    QTabWidget::pane {{
        border: none;
        background: transparent;
        border-radius: {C.RADIUS_MD};
    }}
    QTabBar {{
        background: transparent;
    }}
    QTabBar::tab {{
        background: transparent;
        color: {C.TEXT_SECONDARY};
        font-size: 14px;
        font-weight: 500;
        padding: 10px 20px;
        margin-right: 2px;
        border: none;
        border-radius: {C.RADIUS_MD} {C.RADIUS_MD} 0 0;
        border-bottom: 2px solid transparent;
    }}
    QTabBar::tab:hover {{
        color: {C.TEXT_PRIMARY};
        background: {C.FILL1};
    }}
    QTabBar::tab:selected {{
        color: {C.PRIMARY};
        background: {C.BG_CARD};
        border-bottom: 2px solid {C.PRIMARY};
        font-weight: 600;
    }}

    /* ── 主按钮 (Primary) ── */
    QPushButton#primaryBtn {{
        background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3370FF, stop:1 #165DFF);
        color: #FFFFFF;
        border: none;
        border-radius: 8px;
        padding: 10px 28px;
        font-size: 14px;
        font-weight: 600;
        letter-spacing: 1px;
    }}
    QPushButton#primaryBtn:hover {{
        background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #4080FF, stop:1 #2666FF);
    }}
    QPushButton#primaryBtn:pressed {{
        background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #0E42D2, stop:1 #165DFF);
    }}
    QPushButton#primaryBtn:disabled {{
        background: {C.FILL3};
        color: {C.TEXT_DISABLED};
    }}

    /* ── 次按钮 (Outline) ── */
    QPushButton#outlineBtn {{
        background: {C.BG_CARD};
        color: {C.TEXT_REGULAR};
        border: 1px solid {C.BORDER};
        border-radius: 8px;
        padding: 8px 18px;
        font-size: 13px;
        font-weight: 500;
    }}
    QPushButton#outlineBtn:hover {{
        color: {C.PRIMARY};
        border-color: #A9C7FF;
        background: {C.PRIMARY_LIGHT};
    }}
    QPushButton#outlineBtn:pressed {{
        color: {C.PRIMARY_ACTIVE};
        border-color: {C.PRIMARY};
        background: #D6E4FF;
    }}
    QPushButton#outlineBtn:disabled {{
        color: {C.TEXT_DISABLED};
        border-color: {C.BORDER_LIGHT};
        background: {C.FILL1};
    }}

    /* ── 文字按钮 (Text) ── */
    QPushButton#textBtn {{
        background: transparent;
        color: {C.TEXT_REGULAR};
        border: none;
        border-radius: {C.RADIUS_SM};
        padding: 8px 12px;
        font-size: 13px;
    }}
    QPushButton#textBtn:hover {{
        background: {C.FILL1};
        color: {C.PRIMARY};
    }}
    QPushButton#textBtn:disabled {{
        color: {C.TEXT_DISABLED};
    }}

    /* ── 普通按钮 (兜底) ── */
    QPushButton {{
        background: {C.BG_CARD};
        color: {C.TEXT_REGULAR};
        border: 1px solid {C.BORDER};
        border-radius: 8px;
        padding: 8px 18px;
        font-size: 13px;
        font-weight: 500;
    }}
    QPushButton:hover {{
        color: {C.PRIMARY};
        border-color: #A9C7FF;
        background: {C.PRIMARY_LIGHT};
    }}
    QPushButton:pressed {{
        color: {C.PRIMARY_ACTIVE};
        border-color: {C.PRIMARY};
    }}
    QPushButton:disabled {{
        color: {C.TEXT_DISABLED};
        border-color: {C.BORDER_LIGHT};
        background: {C.FILL1};
    }}

    /* ── 进度条 ── */
    QProgressBar {{
        background: {C.FILL2};
        border: none;
        border-radius: 100px;
        height: 8px;
        text-align: center;
    }}
    QProgressBar::chunk {{
        background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
            stop:0 {C.PRIMARY}, stop:1 {C.PRIMARY_HOVER});
        border-radius: 100px;
    }}

    /* ── 单选按钮 ── */
    QRadioButton {{
        font-size: 14px;
        color: {C.TEXT_PRIMARY};
        spacing: 8px;
        padding: 6px 0;
    }}
    QRadioButton::indicator {{
        width: 16px;
        height: 16px;
        border: 2px solid {C.BORDER};
        border-radius: 50%;
        background: {C.BG_CARD};
    }}
    QRadioButton::indicator:hover {{
        border-color: {C.PRIMARY_HOVER};
    }}
    QRadioButton::indicator:checked {{
        border: 5px solid {C.PRIMARY};
        background: transparent;
    }}

    /* ── 表格 ── */
    QTableWidget {{
        background: {C.BG_CARD};
        alternate-background-color: {C.FILL1};
        border: 1px solid {C.BORDER_LIGHT};
        border-radius: {C.RADIUS_MD};
        font-size: 13px;
        gridline-color: {C.BORDER_LIGHT};
        selection-background-color: {C.PRIMARY_LIGHT};
        selection-color: {C.PRIMARY};
    }}
    QTableWidget::item {{
        padding: 10px 14px;
        border-bottom: 1px solid {C.BORDER_LIGHT};
    }}
    QTableWidget::item:hover {{
        background: {C.FILL1};
    }}
    QHeaderView::section {{
        background: {C.FILL1};
        color: {C.TEXT_REGULAR};
        padding: 10px 14px;
        border: none;
        border-bottom: 1px solid {C.BORDER};
        border-right: 1px solid {C.BORDER_LIGHT};
        font-weight: 600;
        font-size: 13px;
    }}
    QTableWidget QTableWidget {{
        border-radius: 0px;
    }}

    /* ── 消息对话框 (QMessageBox) ── */
    QMessageBox {{
        background-color: {C.BG_CARD};
    }}
    QMessageBox QLabel {{
        color: {C.TEXT_PRIMARY};
        font-size: 14px;
        font-weight: 400;
        background: transparent;
        border: none;
        padding: 0;
        margin: 0;
        min-width: 320px;
        min-height: 24px;
    }}
    QMessageBox QLabel#qt_msgbox_label {{
        color: {C.TEXT_PRIMARY};
        font-size: 14px;
        background: transparent;
    }}
    QMessageBox QPushButton {{
        background: {C.PRIMARY};
        color: #FFFFFF;
        border: none;
        border-radius: 6px;
        padding: 8px 24px;
        font-size: 14px;
        font-weight: 600;
        min-width: 80px;
        min-height: 32px;
    }}
    QMessageBox QPushButton:hover {{
        background: {C.PRIMARY_HOVER};
    }}
    QMessageBox QPushButton:pressed {{
        background: {C.PRIMARY_ACTIVE};
    }}

    /* ── 对话框 (QDialog) 兜底 ── */
    QDialog {{
        background-color: {C.BG_CARD};
        color: {C.TEXT_PRIMARY};
    }}
    QDialog QLabel {{
        color: {C.TEXT_PRIMARY};
        background: transparent;
    }}
    QDialog QTextEdit {{
        color: {C.TEXT_PRIMARY};
    }}
    QDialog QLineEdit {{
        color: {C.TEXT_PRIMARY};
    }}
    QDialog QComboBox {{
        color: {C.TEXT_PRIMARY};
    }}

    /* ── 文件对话框 ── */
    QFileDialog {{
        background-color: {C.BG_CARD};
        color: {C.TEXT_PRIMARY};
    }}
    QFileDialog QLabel {{
        color: {C.TEXT_PRIMARY};
    }}
    QFileDialog QLineEdit {{
        color: {C.TEXT_PRIMARY};
    }}
    QFileDialog QComboBox {{
        color: {C.TEXT_PRIMARY};
    }}
    QFileDialog QTreeView {{
        color: {C.TEXT_PRIMARY};
        background: {C.BG_CARD};
    }}
    QFileDialog QListView {{
        color: {C.TEXT_PRIMARY};
        background: {C.BG_CARD};
    }}

    /* ── 滚动条 ── */
    QScrollBar:vertical {{
        background: transparent;
        width: 6px;
        margin: 0;
    }}
    QScrollBar::handle:vertical {{
        background: {C.FILL4};
        border-radius: 3px;
        min-height: 20px;
    }}
    QScrollBar::handle:vertical:hover {{
        background: {C.TEXT_SECONDARY};
    }}
    QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
        height: 0;
    }}
    QScrollBar:horizontal {{
        background: transparent;
        height: 6px;
        margin: 0;
    }}
    QScrollBar::handle:horizontal {{
        background: {C.FILL4};
        border-radius: 3px;
        min-width: 20px;
    }}
    QScrollBar::handle:horizontal:hover {{
        background: {C.TEXT_SECONDARY};
    }}
    QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
        width: 0;
    }}
    """


# ═══════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════
def make_h_line(color=None):
    """水平分隔线"""
    C = ArcoColors
    line = QFrame()
    line.setFrameShape(QFrame.Shape.HLine)
    line.setFixedHeight(1)
    line.setStyleSheet(f"background-color: {color or C.BORDER_LIGHT}; border: none; margin: 4px 0;")
    return line


def make_shadow_widget(widget, shadow_type="sm"):
    """给 widget 添加 Arco Design 风格阴影"""
    shadow = QGraphicsDropShadowEffect()
    shadow.setBlurRadius(12)
    shadow.setOffset(0, 2)
    shadow.setColor(QColor(0, 0, 0, 20))
    if shadow_type == "md":
        shadow.setBlurRadius(20)
        shadow.setOffset(0, 4)
        shadow.setColor(QColor(0, 0, 0, 30))
    widget.setGraphicsEffect(shadow)


# ═══════════════════════════════════════════════════════════
# AI 调用
# ═══════════════════════════════════════════════════════════
def chat(system, user, temperature=0.3):
    if not _api_key or _api_key == "sk-placeholder":
        raise RuntimeError(
            "OPENAI_API_KEY 未配置。请在 .env 文件中设置 OPENAI_API_KEY，"
            "或将 .env 文件放在程序同级目录下。"
        )
    logger.info(f"调用模型: {model}, temperature={temperature}")
    resp = client.chat.completions.create(
        model=model, temperature=temperature,
        messages=[{"role": "system", "content": system},
                  {"role": "user", "content": user}]
    )
    return resp.choices[0].message.content


def fetch_url(url):
    try:
        resp = requests.get(f"https://r.jina.ai/{url}", timeout=15)
        if resp.status_code == 200 and len(resp.text) > 200:
            return resp.text
    except Exception:
        pass
    try:
        resp = requests.get(url, timeout=15)
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        text = soup.get_text(separator="\n").strip()
        if len(text) > 200:
            return text
    except Exception:
        pass
    raise RuntimeError("无法抓取URL，请手动复制正文。")


def split_chunks(text, max_chars=8000):
    """智能分块，兼容中文（按字符数）和英文（按空格分词）。"""
    # 判断是否以 CJK 字符为主（超过 30%）
    cjk_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff' or '\u3040' <= c <= '\u30ff' or '\uac00' <= c <= '\ud7af')
    is_cjk = cjk_count > len(text) * 0.3

    if is_cjk:
        # 中文：按字符数分割，优先在句号/段落边界断开
        if len(text) <= max_chars:
            return [text]
        break_chars = '。！？\n；'
        chunks, current = [], []
        for ch in text:
            current.append(ch)
            if len(current) >= max_chars:
                # 向回找到最近的断句位置
                best = -1
                for j in range(len(current) - 1, max(0, len(current) - 200), -1):
                    if current[j] in break_chars:
                        best = j + 1
                        break
                if best > 0:
                    chunks.append("".join(current[:best]))
                    current = current[best:]
                else:
                    chunks.append("".join(current))
                    current = []
        if current:
            chunks.append("".join(current))
        return chunks if chunks else [text]
    else:
        # 英文/西文：按空格分词
        words = text.split()
        if len(words) <= 3000:
            return [text]
        chunks, current = [], []
        for word in words:
            current.append(word)
            if len(current) >= 3000:
                chunks.append(" ".join(current))
                current = []
        if current:
            chunks.append(" ".join(current))
        return chunks


# ═══════════════════════════════════════════════════════════
# 五步翻译工作流
# ═══════════════════════════════════════════════════════════
def step1_analyze(text, source_lang, target_lang, audience, style):
    """深度分析源文本，长文档分段分析后合并术语，确保覆盖完整文档。"""
    ANALYSIS_CHUNK = 6000

    sys_prompt = """你是专业翻译前置分析师。
输出严格按照以下固定Markdown结构：
## 概要
## 术语表（原文 → 译法）
请逐行列出所有专业术语、缩写、惯用表达、专有名词。
格式要求：每行一条，用 → 连接，例如：
- machine learning → 机器学习
- large language model → 大语言模型
- the state of the art → 最先进的
## 语气与风格判定
## 读者理解难点 & 文化注解点
## 修辞隐喻与替换映射
"""

    if len(text) <= ANALYSIS_CHUNK:
        # 短文档：一次分析
        user_prompt = f"""
源语言：{source_lang}
目标语言：{target_lang}
目标读者：{audience or "普通读者"}
风格模式：{style or "auto"}

原文：
{text}
"""
        return chat(sys_prompt, user_prompt, temperature=0.2)

    # ── 长文档：分段分析 ──
    # 第一段：完整分析（概要 + 术语 + 风格 + 难点 + 修辞）
    chunks = [text[i:i+ANALYSIS_CHUNK] for i in range(0, len(text), ANALYSIS_CHUNK)]
    user_prompt = f"""
源语言：{source_lang}
目标语言：{target_lang}
目标读者：{audience or "普通读者"}
风格模式：{style or "auto"}
这是长文档的第 1/{len(chunks)} 段，请进行完整分析。

原文：
{chunks[0]}
"""
    full_analysis = chat(sys_prompt, user_prompt, temperature=0.2)

    # 后续段：只提取术语（避免重复生成概要/风格）
    term_prompt = """你是术语提取专家。从以下文本中提取所有专业术语、缩写、惯用表达、专有名词。
只输出术语列表，格式：- 原文 → 译文
不要输出其他内容。"""

    for idx in range(1, len(chunks)):
        chunk_terms = chat(
            term_prompt,
            f"源语言：{source_lang}，目标语言：{target_lang}\n\n原文：\n{chunks[idx]}",
            temperature=0.1
        )
        # 将术语追加到分析报告的术语表部分
        if chunk_terms.strip():
            full_analysis += f"\n\n### 第 {idx+1}/{len(chunks)} 段补充术语\n{chunk_terms.strip()}"

    return full_analysis


def step2_build_prompt(analysis, source_lang, target_lang, audience, style):
    from glossary import get_glossary_prompt_block

    if style == "auto" or not style.strip():
        style_block = f"""The source text's tone is extracted from analysis above.
Reproduce this tone faithfully in {target_lang}. Match register, rhythm, personality exactly."""
    else:
        style_block = f"""{style}
Apply this style consistently across full text."""

    glossary_block = get_glossary_prompt_block(source_lang, target_lang)

    prompt_full = f"""You are a professional translator. Translate from {source_lang} to {target_lang}.
Think and reason entirely in {target_lang}.

## Target Audience
{audience or "general readers"}

## Translation Style
{style_block}

{glossary_block}

## Content Background
{analysis}

## Translation Principles
- Complete every sentence, no omission / summarization
- Accuracy first, meaning over literal word
- Keep markdown format fully preserved
- Cultural terms add brief explanation in parentheses
- Natural target language expression, adjust sentence structure freely
"""
    return prompt_full


def step3_draft(source_text, prompt, source_lang, target_lang):
    chunks = split_chunks(source_text)
    subagent_cmd = """Follow all rules in the translation context above.
Translate this chunk COMPLETELY, every sentence, no omission, no summarization.
Keep paragraph count similar, preserve format fully. Only output pure translation result.
"""
    if len(chunks) == 1:
        return chat(prompt + "\n" + subagent_cmd, source_text)
    results = []
    for chunk in chunks:
        t = chat(prompt + "\n" + subagent_cmd, chunk)
        results.append(t)
    return "\n\n".join(results)


def step4_critique(source_text, draft, analysis, source_lang, target_lang):
    """审校译文 — 长文档分段审校，避免截断。"""
    # 对长文档进行分段审校
    max_len = 6000
    src_chunks = [source_text[i:i+max_len] for i in range(0, len(source_text), max_len)]
    draft_chunks = [draft[i:i+max_len] for i in range(0, len(draft), max_len)]
    # 合并对应段进行审校
    pairs = []
    for i in range(max(len(src_chunks), len(draft_chunks))):
        s = src_chunks[i] if i < len(src_chunks) else ""
        d = draft_chunks[i] if i < len(draft_chunks) else ""
        if s or d:
            pairs.append((s, d))

    critiques = []
    for idx, (src_part, draft_part) in enumerate(pairs):
        part_label = f"（第 {idx+1}/{len(pairs)} 段）" if len(pairs) > 1 else ""
        critique = chat(
            "你是严格的翻译审校专家。",
            f"""审校{source_lang}→{target_lang}译文{part_label}，输出诊断报告：
## 准确性与完整性
## 翻译腔问题（原句 → 建议）
## 修辞与情感保真
## 表达与逻辑
## 总结
只列出问题，不要自行改写译文原文。

原文：\n{src_part}
译文：\n{draft_part}
分析要点：\n{analysis[:2000]}"""
        )
        critiques.append(critique)

    return "\n\n---\n\n".join(critiques)


def step5_final(draft, critique, target_lang):
    """终稿润色 — 长文档分段修正，避免超出上下文窗口。"""
    max_len = 6000
    draft_chunks = [draft[i:i+max_len] for i in range(0, len(draft), max_len)]

    if len(draft_chunks) <= 1:
        return chat(
            f"你是{target_lang}母语精修专家，严格依据审校报告逐条修正",
            f"""原文分析、初译草稿、审校问题全部参考上方。
逐条修正术语、表达、不通顺、翻译腔；保持原意不变。
只输出最终纯净定稿译文：
初译：{draft}
审校：{critique}"""
        )

    # 长文档：逐段修正
    # 将 critique 也分段，确保对应关系
    critique_chunks = [critique[i:i+max_len] for i in range(0, len(critique), max_len)]
    final_parts = []
    for idx, draft_part in enumerate(draft_chunks):
        # 取对应的 critique 段（可能不一一对应，取最后一段作为兜底）
        crit_part = critique_chunks[idx] if idx < len(critique_chunks) else critique_chunks[-1] if critique_chunks else ""
        result = chat(
            f"你是{target_lang}母语精修专家，严格依据审校报告逐条修正",
            f"""以下是第 {idx+1}/{len(draft_chunks)} 段的修正任务。
逐条修正术语、表达、不通顺、翻译腔；保持原意不变。
只输出这一段的最终纯净定稿译文：
初译段落：{draft_part}
审校参考：{crit_part}"""
        )
        final_parts.append(result)

    return "\n\n".join(final_parts)


# ═══════════════════════════════════════════════════════════
# 字幕翻译线程
# ═══════════════════════════════════════════════════════════
class SubtitleWorker(QThread):
    """字幕翻译：分析 → 初译 → 审校 → 终稿，保留时间轴"""
    progress = pyqtSignal(str, int)
    analysis_done = pyqtSignal(str)          # AI 分析报告
    critique_done = pyqtSignal(str)          # AI 审校报告
    finished = pyqtSignal(object, list)      # SubtitleFile, translated_texts
    error = pyqtSignal(str)

    def __init__(self, subtitle, source_lang, target_lang,
                 style="", audience=""):
        super().__init__()
        self.subtitle = subtitle
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.style = style or ""
        self.audience = audience or ""

    # ── 第一步：AI 分析字幕内容 ──
    def _step_analyze(self, source_text):
        """分段分析字幕内容，长字幕确保覆盖完整。"""
        ANALYSIS_CHUNK = 8000

        sys_prompt = """你是专业翻译前置分析师，当前任务是对字幕文本进行分析。
字幕文本每行是一条字幕，格式如 [001] Hello world。
输出严格按照以下固定Markdown结构：
## 概要
## 术语表（原文 → 译法）
请逐行列出所有专业术语、缩写、惯用表达、专有名词。
格式要求：每行一条，用 → 连接，例如：
- machine learning → 机器学习
- large language model → 大语言模型
## 语气与风格判定
## 读者理解难点 & 文化注解点
## 修辞隐喻与替换映射
"""

        if len(source_text) <= ANALYSIS_CHUNK:
            # 短字幕：一次分析
            user_prompt = f"""
源语言：{self.source_lang}
目标语言：{self.target_lang}
目标读者：{self.audience or "普通观众"}
风格模式：{self.style or "auto"}
这是字幕文件，每行代表一条字幕，请基于字幕文本内容进行分析。

原文：
{source_text}
"""
            return chat(sys_prompt, user_prompt, temperature=0.2)

        # ── 长字幕：分段分析 ──
        chunks = [source_text[i:i+ANALYSIS_CHUNK] for i in range(0, len(source_text), ANALYSIS_CHUNK)]

        # 第一段：完整分析
        user_prompt = f"""
源语言：{self.source_lang}
目标语言：{self.target_lang}
目标读者：{self.audience or "普通观众"}
风格模式：{self.style or "auto"}
这是字幕文件，每行代表一条字幕。
这是长字幕的第 1/{len(chunks)} 段，请进行完整分析。

原文：
{chunks[0]}
"""
        full_analysis = chat(sys_prompt, user_prompt, temperature=0.2)

        # 后续段：只提取术语
        term_prompt = """你是术语提取专家。从以下字幕文本中提取所有专业术语、缩写、惯用表达、专有名词。
只输出术语列表，格式：- 原文 → 译文
不要输出其他内容。"""

        for idx in range(1, len(chunks)):
            self.progress.emit(
                f"分析字幕内容 {idx+1}/{len(chunks)}...", 1)
            chunk_terms = chat(
                term_prompt,
                f"源语言：{self.source_lang}，目标语言：{self.target_lang}\n\n原文：\n{chunks[idx]}",
                temperature=0.1
            )
            if chunk_terms.strip():
                full_analysis += f"\n\n### 第 {idx+1}/{len(chunks)} 段补充术语\n{chunk_terms.strip()}"

        return full_analysis

    # ── 第二步：组装翻译提示 ──
    def _step_build_prompt(self, analysis):
        from glossary import get_glossary_prompt_block

        if self.style == "auto" or not self.style.strip():
            style_block = f"""The source text's tone is extracted from analysis above.
Reproduce this tone faithfully in {self.target_lang}. Match register, rhythm, personality exactly.
字幕翻译要求：每行一条，简洁自然，适合屏幕阅读。"""
        else:
            style_block = f"""{self.style}
Apply this style consistently across full text.
字幕翻译要求：每行一条，简洁自然，适合屏幕阅读。"""

        glossary_block = get_glossary_prompt_block(self.source_lang, self.target_lang)

        prompt_full = f"""You are a professional subtitle translator translating from {self.source_lang} to {self.target_lang}.
Think and reason entirely in {self.target_lang}.

## Target Audience
{self.audience or "general viewers"}

## Translation Style
{style_block}

{glossary_block}

## Content Background
{analysis}

## Subtitle Translation Principles
- Translate EACH line independently, output EXACTLY the same number of lines
- Keep translations concise and natural — suitable for on-screen subtitle reading
- No omission, no summarization
- Cultural terms add brief explanation in parentheses
- Do NOT add explanations, notes, or comments
- Do NOT merge or split lines
"""
        return prompt_full

    # ── 第三步：初译 ──
    def _step_translate(self, prompt):
        from subtitle_handler import (
            parse_subtitle_translation_response,
            _apply_translations,
        )
        batch_size = 50
        all_translated = []
        entries_with_text = [e for e in self.subtitle.entries if e.text.strip()]
        total_batches = (len(entries_with_text) + batch_size - 1) // batch_size

        for batch_idx in range(total_batches):
            start = batch_idx * batch_size
            end = min(start + batch_size, len(entries_with_text))
            batch_entries = entries_with_text[start:end]

            self.progress.emit(
                f"初译字幕 {start+1}-{end}/{len(entries_with_text)}...",
                3 + int(1 * batch_idx / max(total_batches, 1)))

            batch_input = "\n".join(
                f"[{i+1:03d}] {e.text}" for i, e in enumerate(batch_entries))

            sub_cmd = """Follow all rules above. Translate each numbered line COMPLETELY.
Keep the [NNN] prefix in your response. Only output translated lines, nothing else."""
            response = chat(prompt + "\n" + sub_cmd, batch_input, temperature=0.3)
            batch_translated = parse_subtitle_translation_response(
                response, len(batch_entries))
            all_translated.extend(batch_translated)

        _apply_translations(self.subtitle, all_translated)
        return all_translated

    # ── 第四步：审校 ──
    def _step_critique(self, source_text, draft_text, analysis):
        max_len = 6000
        src_chunks = [source_text[i:i+max_len] for i in range(0, len(source_text), max_len)]
        draft_chunks = [draft_text[i:i+max_len] for i in range(0, len(draft_text), max_len)]
        pairs = []
        for i in range(max(len(src_chunks), len(draft_chunks))):
            s = src_chunks[i] if i < len(src_chunks) else ""
            d = draft_chunks[i] if i < len(draft_chunks) else ""
            if s or d:
                pairs.append((s, d))

        critiques = []
        for idx, (src_part, draft_part) in enumerate(pairs):
            part_label = f"（第 {idx+1}/{len(pairs)} 段）" if len(pairs) > 1 else ""
            critique = chat(
                "你是严格的翻译审校专家，当前审校对象是字幕翻译。",
                f"""审校{self.source_lang}→{self.target_lang}字幕译文{part_label}，输出诊断报告：
## 准确性与完整性
## 翻译腔问题（原句 → 建议）
## 修辞与情感保真
## 表达与逻辑
## 总结
只列出问题，不要自行改写译文原文。

原文：\n{src_part}
译文：\n{draft_part}
分析要点：\n{analysis[:2000]}"""
            )
            critiques.append(critique)

        return "\n\n---\n\n".join(critiques)

    # ── 第五步：终稿润色 ──
    def _step_final(self, draft_text, critique):
        max_len = 6000
        draft_chunks = [draft_text[i:i+max_len] for i in range(0, len(draft_text), max_len)]

        if len(draft_chunks) <= 1:
            return chat(
                f"你是{self.target_lang}母语精修专家，当前润色对象是字幕翻译，严格依据审校报告逐条修正",
                f"""逐条修正术语、表达、不通顺、翻译腔；保持原意不变。
只输出最终纯净定稿译文（保持 [NNN] 编号格式）：
初译：{draft_text}
审校：{critique}"""
            )

        critique_chunks = [critique[i:i+max_len] for i in range(0, len(critique), max_len)]
        final_parts = []
        for idx, draft_part in enumerate(draft_chunks):
            crit_part = critique_chunks[idx] if idx < len(critique_chunks) else critique_chunks[-1] if critique_chunks else ""
            result = chat(
                f"你是{self.target_lang}母语精修专家，当前润色对象是字幕翻译，严格依据审校报告逐条修正",
                f"""以下是第 {idx+1}/{len(draft_chunks)} 段的修正任务。
逐条修正术语、表达、不通顺、翻译腔；保持原意不变。
只输出这一段的最终纯净定稿译文（保持 [NNN] 编号格式）：
初译段落：{draft_part}
审校参考：{crit_part}"""
            )
            final_parts.append(result)

        return "\n".join(final_parts)

    # ── 主流程 ──
    def run(self):
        try:
            from subtitle_handler import (
                parse_subtitle_translation_response,
                _apply_translations,
            )
            from glossary import extract_terms_from_analysis, add_terms_batch

            entries_with_text = [e for e in self.subtitle.entries if e.text.strip()]
            source_text = "\n".join(
                f"[{i+1:03d}] {e.text}" for i, e in enumerate(entries_with_text))

            # ── 第一步：深度分析 ──
            self.progress.emit("第一步：深度分析字幕内容...", 1)
            analysis = self._step_analyze(source_text)
            self.analysis_done.emit(analysis)

            # 术语自动入库
            new_terms = extract_terms_from_analysis(analysis, self.source_lang, self.target_lang)
            if new_terms:
                added, _ = add_terms_batch(new_terms, self.source_lang, self.target_lang)
                self.progress.emit(f"术语入库：新增 {added} 条", 1)

            # ── 第二步：组装提示 ──
            self.progress.emit("第二步：组装翻译提示...", 2)
            prompt = self._step_build_prompt(analysis)

            # ── 第三步：初译 ──
            self.progress.emit("第三步：初译字幕...", 3)
            all_translated = self._step_translate(prompt)

            # 构建完整译文文本（用于审校）
            draft_text = "\n".join(
                f"[{i+1:03d}] {t}" for i, t in enumerate(all_translated))

            # ── 第四步：审校 ──
            self.progress.emit("第四步：审校译文...", 4)
            critique = self._step_critique(source_text, draft_text, analysis)
            self.critique_done.emit(critique)

            # ── 第五步：终稿润色 ──
            self.progress.emit("第五步：终稿润色...", 5)
            final_text = self._step_final(draft_text, critique)

            # 将终稿结果重新应用到字幕对象
            final_translated = parse_subtitle_translation_response(
                final_text, len(entries_with_text))
            _apply_translations(self.subtitle, final_translated)

            self.progress.emit("字幕翻译完成！", 5)
            self.finished.emit(self.subtitle, final_translated)

        except Exception as e:
            logger.exception("字幕翻译异常")
            self.error.emit(str(e))


# ═══════════════════════════════════════════════════════════
# 翻译线程
# ═══════════════════════════════════════════════════════════
class TranslateWorker(QThread):
    progress = pyqtSignal(str, int)
    analysis_done = pyqtSignal(str)
    critique_done = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, source_text, source_lang, target_lang,
                 style, audience):
        super().__init__()
        self.source_text = source_text
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.style = style
        self.audience = audience

    def run(self):
        try:
            self.progress.emit("第一步：深度分析...", 1)
            analysis = step1_analyze(self.source_text, self.source_lang,
                                     self.target_lang, self.audience, self.style)
            self.analysis_done.emit(analysis)

            from glossary import extract_terms_from_analysis, add_terms_batch
            new_terms = extract_terms_from_analysis(analysis, self.source_lang, self.target_lang)
            if new_terms:
                added, _ = add_terms_batch(new_terms, self.source_lang, self.target_lang)
                self.progress.emit(f"术语入库：新增 {added} 条", 1)

            self.progress.emit("第二步：组装提示...", 2)
            prompt = step2_build_prompt(analysis, self.source_lang,
                                        self.target_lang, self.audience, self.style)

            self.progress.emit("第三步：初译...", 3)
            draft = step3_draft(self.source_text, prompt,
                                self.source_lang, self.target_lang)

            self.progress.emit("第四步：审校...", 4)
            critique = step4_critique(self.source_text, draft, analysis,
                                      self.source_lang, self.target_lang)
            self.critique_done.emit(critique)

            self.progress.emit("第五步：终稿润色...", 5)
            final = step5_final(draft, critique, self.target_lang)
            self.finished.emit(final)

        except Exception as e:
            self.error.emit(str(e))


# ═══════════════════════════════════════════════════════════
# 导出对话框 — Arco Design 风格
# ═══════════════════════════════════════════════════════════
class ExportDialog(QDialog):
    def __init__(self, parent, file_type):
        super().__init__(parent)
        self.setWindowTitle("导出设置")
        self.setMinimumWidth(360)
        self.setMinimumHeight(420)
        C = ArcoColors
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {C.BG_PAGE};
                color: {C.TEXT_PRIMARY};
            }}
            QLabel {{
                font-size: 14px;
                color: {C.TEXT_PRIMARY};
                font-weight: 600;
            }}
            QRadioButton {{
                font-size: 14px;
                color: {C.TEXT_PRIMARY};
                spacing: 8px;
                padding: 7px 12px;
                border-radius: {C.RADIUS_MD};
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
            }}
            QRadioButton:hover {{
                border-color: {C.PRIMARY_HOVER};
                background: {C.PRIMARY_LIGHT2};
            }}
            QRadioButton::indicator {{
                width: 16px;
                height: 16px;
                border: 2px solid {C.BORDER};
                border-radius: 50%;
                background: {C.BG_CARD};
            }}
            QRadioButton::indicator:checked {{
                border: 5px solid {C.PRIMARY};
            }}
            QPushButton {{
                background: {C.PRIMARY};
                color: #FFFFFF;
                border: none;
                border-radius: {C.RADIUS_MD};
                padding: 10px 24px;
                font-size: 14px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background: {C.PRIMARY_HOVER};
            }}
        """)

        card = QFrame()
        card.setStyleSheet(f"""
            QFrame {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_LG};
                padding: 24px;
            }}
        """)
        make_shadow_widget(card, "md")
        card_layout = QVBoxLayout(card)
        card_layout.setSpacing(16)
        card_layout.setContentsMargins(24, 24, 24, 24)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(24, 24, 24, 24)
        outer.addWidget(card)

        card_layout.addWidget(QLabel("导出格式"))

        self.fmt_group = QButtonGroup()
        fmt_options = {
            "pdf": "PDF",
            "docx": "Word (.docx)",
            "doc": "Word (.doc)",
            "xlsx": "Excel (.xlsx)",
            "xls": "Excel (.xls)",
            "pptx": "PowerPoint (.pptx)",
            "ppt": "PowerPoint (.ppt)",
        }
        default_fmt = file_type if file_type in fmt_options else "docx"
        fmt_grid = QVBoxLayout()
        fmt_grid.setSpacing(6)
        for fmt, label in fmt_options.items():
            rb = QRadioButton(label)
            rb.setProperty("val", fmt)
            if fmt == default_fmt:
                rb.setChecked(True)
            self.fmt_group.addButton(rb)
            fmt_grid.addWidget(rb)
        card_layout.addLayout(fmt_grid)

        card_layout.addWidget(make_h_line())
        card_layout.addWidget(QLabel("导出模式"))

        self.mode_group = QButtonGroup()
        mode_grid = QVBoxLayout()
        mode_grid.setSpacing(6)
        for mode, label in [
            ("paragraph", "段落对照模式（推荐）"),
            ("bilingual", "行对照模式（原文+译文）"),
            ("translation", "仅译文模式"),
        ]:
            rb = QRadioButton(label)
            rb.setProperty("val", mode)
            if mode == "paragraph":
                rb.setChecked(True)
            self.mode_group.addButton(rb)
            mode_grid.addWidget(rb)
        card_layout.addLayout(mode_grid)

        card_layout.addStretch()

        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn = QPushButton("确认导出")
        btn.setCursor(Qt.CursorShape.PointingHandCursor)
        btn.clicked.connect(self.accept)
        btn_row.addWidget(btn)
        card_layout.addLayout(btn_row)

    def get_fmt(self):
        for b in self.fmt_group.buttons():
            if b.isChecked():
                return b.property("val")

    def get_mode(self):
        for b in self.mode_group.buttons():
            if b.isChecked():
                return b.property("val")


# ═══════════════════════════════════════════════════════════
# 主窗口 — Arco Design 风格
# ═══════════════════════════════════════════════════════════
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Translation Agent")
        self.setMinimumSize(1060, 820)
        self.file_path = None
        self.file_type = None
        self.file_extra = None
        self.last_result = ""
        self.last_source = ""
        self.last_analysis = ""
        self.last_critique = ""
        # 字幕相关状态
        self.subtitle_obj = None
        self.subtitle_path = None
        self.subtitle_worker = None
        self.apply_styles()
        self.setup_ui()

    def apply_styles(self):
        self.setStyleSheet(arco_global_stylesheet())

    def _swap_languages(self):
        """交换源语言和目标语言（按文本匹配，避免因列表顺序不同导致失效）"""
        src_text = self.source_lang.currentText()
        tgt_text = self.target_lang.currentText()
        if src_text == tgt_text:
            return
        # 在目标语言列表中找源语言文本
        src_idx_in_tgt = self.target_lang.findText(src_text)
        tgt_idx_in_src = self.source_lang.findText(tgt_text)
        if tgt_idx_in_src >= 0:
            self.source_lang.setCurrentIndex(tgt_idx_in_src)
        if src_idx_in_tgt >= 0:
            self.target_lang.setCurrentIndex(src_idx_in_tgt)

    def setup_ui(self):
        C = ArcoColors

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(12)
        layout.setContentsMargins(24, 20, 24, 20)

        # ── 顶部标题栏 ──
        header = QFrame()
        header.setStyleSheet("background: transparent; border: none;")
        header_layout = QVBoxLayout(header)
        header_layout.setSpacing(4)
        header_layout.setContentsMargins(0, 0, 0, 0)

        title = QLabel("Translation Agent")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 22, QFont.Weight.Bold))
        title.setStyleSheet(f"color: {C.TEXT_PRIMARY}; background: transparent; border: none;")
        header_layout.addWidget(title)

        sub = QLabel("支持 PDF / Word / Excel / PPT / 字幕(SRT/VTT/ASS)  ·  分析 → 提示 → 初译 → 审校 → 终稿")
        sub.setAlignment(Qt.AlignmentFlag.AlignCenter)
        sub.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; background: transparent; border: none; letter-spacing: 0.5px;")
        header_layout.addWidget(sub)

        layout.addWidget(header)

        # ── 翻译设置卡片 ──
        sg = QFrame()
        sg.setObjectName("card")
        sg.setStyleSheet(f"""
            QFrame#card {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_LG};
            }}
        """)
        make_shadow_widget(sg, "sm")
        sl = QHBoxLayout(sg)
        sl.setSpacing(12)
        sl.setContentsMargins(20, 12, 20, 12)

        # 卡片小标题（放在卡片上方）
        settings_title = QLabel("翻译设置")
        settings_title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 13, QFont.Weight.Bold))
        settings_title.setStyleSheet(f"color: {C.TEXT_PRIMARY}; border: none; background: transparent;")

        layout.addWidget(settings_title)
        layout.addWidget(sg)

        # ── 语言选择区（紧凑单行） ──
        lang_row = QHBoxLayout()
        lang_row.setSpacing(8)

        # 源语言
        src_lbl = QLabel("源语言")
        src_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        lang_row.addWidget(src_lbl)
        self.source_lang = QComboBox()
        self.source_lang.addItems(["English", "Chinese", "Japanese", "Korean", "French", "German", "Spanish"])
        self.source_lang.setFixedHeight(34)
        self.source_lang.setMinimumWidth(120)
        self.source_lang.setCursor(Qt.CursorShape.PointingHandCursor)
        lang_row.addWidget(self.source_lang)

        # 交换按钮
        swap_btn = QPushButton("⇄")
        swap_btn.setObjectName("swapBtn")
        swap_btn.setFixedSize(40, 34)
        swap_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        swap_btn.setToolTip("交换源语言与目标语言")
        swap_btn.setStyleSheet(f"""
            QPushButton#swapBtn {{
                background: {C.PRIMARY};
                color: #FFFFFF;
                border: none;
                border-radius: 6px;
                font-size: 16px;
                font-weight: 700;
            }}
            QPushButton#swapBtn:hover {{
                background: {C.PRIMARY_HOVER};
            }}
            QPushButton#swapBtn:pressed {{
                background: {C.PRIMARY_ACTIVE};
            }}
        """)
        swap_btn.clicked.connect(self._swap_languages)
        lang_row.addWidget(swap_btn)

        # 目标语言
        tgt_lbl = QLabel("目标语言")
        tgt_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        lang_row.addWidget(tgt_lbl)
        self.target_lang = QComboBox()
        self.target_lang.addItems(["Chinese", "English", "Japanese", "Korean", "French", "German", "Spanish"])
        self.target_lang.setFixedHeight(34)
        self.target_lang.setMinimumWidth(120)
        self.target_lang.setCursor(Qt.CursorShape.PointingHandCursor)
        lang_row.addWidget(self.target_lang)

        # 分隔点
        sep_dot = QLabel("·")
        sep_dot.setStyleSheet(f"color: {C.TEXT_DISABLED}; font-size: 18px; border: none; background: transparent;")
        sep_dot.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lang_row.addWidget(sep_dot)

        # 风格
        style_lbl = QLabel("风格")
        style_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        lang_row.addWidget(style_lbl)
        self.style_input = QLineEdit()
        self.style_input.setPlaceholderText("formal / auto")
        self.style_input.setFixedHeight(34)
        self.style_input.setMinimumWidth(140)
        lang_row.addWidget(self.style_input)

        # 目标读者
        aud_lbl = QLabel("目标读者")
        aud_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        lang_row.addWidget(aud_lbl)
        self.audience_input = QLineEdit()
        self.audience_input.setPlaceholderText("general / technical")
        self.audience_input.setFixedHeight(34)
        self.audience_input.setMinimumWidth(140)
        lang_row.addWidget(self.audience_input)

        lang_row.addStretch()
        sl.addLayout(lang_row)

        # ── 输入内容卡片 ──
        ig = QFrame()
        ig.setObjectName("inputCard")
        ig.setStyleSheet(f"""
            QFrame#inputCard {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_LG};
            }}
        """)
        make_shadow_widget(ig, "sm")
        il = QVBoxLayout(ig)
        il.setSpacing(8)
        il.setContentsMargins(20, 16, 20, 16)

        input_title = QLabel("输入内容")
        input_title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 13, QFont.Weight.Bold))
        input_title.setStyleSheet(f"color: {C.TEXT_PRIMARY}; border: none; background: transparent;")
        il.addWidget(input_title)

        self.tabs = QTabWidget()
        self.tabs.setStyleSheet(f"""
            QTabWidget::pane {{
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_MD};
                background: {C.BG_CARD};
                padding: 4px;
            }}
        """)

        # Tab 1: 粘贴文字
        tw = QWidget()
        tl = QVBoxLayout(tw)
        tl.setContentsMargins(8, 8, 8, 8)
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("在此粘贴需要翻译的文字...")
        self.text_input.setMinimumHeight(160)
        from PyQt6.QtGui import QTextOption
        self.text_input.setWordWrapMode(QTextOption.WrapMode.WordWrap)
        tl.addWidget(self.text_input)
        self.tabs.addTab(tw, "  粘贴文字  ")

        # Tab 2: URL
        uw = QWidget()
        ul = QVBoxLayout(uw)
        ul.setContentsMargins(8, 12, 8, 8)
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("https://example.com/article")
        self.url_input.setFixedHeight(38)
        ul.addWidget(self.url_input)
        ul.addStretch()
        self.tabs.addTab(uw, "  URL  ")

        # Tab 3: 上传文件
        fw = QWidget()
        fl = QVBoxLayout(fw)
        fl.setContentsMargins(8, 12, 8, 8)
        fbl = QHBoxLayout()
        self.file_btn = QPushButton("选择文件")
        self.file_btn.setObjectName("outlineBtn")
        self.file_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.file_btn.clicked.connect(self.choose_file)
        self.file_label = QLabel("支持 .pdf / .docx / .doc / .xlsx / .xls / .pptx / .ppt")
        self.file_label.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; border: none; background: transparent;")
        fbl.addWidget(self.file_btn)
        fbl.addWidget(self.file_label)
        fbl.addStretch()
        fl.addLayout(fbl)
        fl.addStretch()
        self.tabs.addTab(fw, "  上传文件  ")

        # Tab 4: 字幕翻译
        sw = QWidget()
        sl_layout = QVBoxLayout(sw)
        sl_layout.setContentsMargins(8, 12, 8, 8)

        sub_btn_row = QHBoxLayout()
        self.sub_file_btn = QPushButton("导入字幕")
        self.sub_file_btn.setObjectName("outlineBtn")
        self.sub_file_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.sub_file_btn.clicked.connect(self.choose_subtitle_file)
        self.sub_file_label = QLabel("支持 .srt / .vtt / .ass / .ssa")
        self.sub_file_label.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; border: none; background: transparent;")
        sub_btn_row.addWidget(self.sub_file_btn)
        sub_btn_row.addWidget(self.sub_file_label)
        sub_btn_row.addStretch()
        sl_layout.addLayout(sub_btn_row)

        # 输出模式选择行
        mode_row = QHBoxLayout()
        mode_row.setSpacing(12)
        mode_lbl = QLabel("输出模式：")
        mode_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        mode_row.addWidget(mode_lbl)
        self.sub_mode_combo = QComboBox()
        self.sub_mode_combo.setMinimumWidth(240)
        self.sub_mode_combo.setFixedHeight(34)
        self.sub_mode_combo.addItems([
            "双语字幕（原文+译文，带时间轴）",
            "纯译文（仅译文，带时间轴）",
            "Clean 纯文本（仅译文，无时间轴）",
            "Clean 双语（原文+译文，无时间轴）",
        ])
        mode_row.addWidget(self.sub_mode_combo)
        fmt_lbl = QLabel("输出格式：")
        fmt_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; font-weight: 500; border: none; background: transparent;")
        mode_row.addWidget(fmt_lbl)
        self.sub_fmt_combo = QComboBox()
        self.sub_fmt_combo.setMinimumWidth(120)
        self.sub_fmt_combo.setFixedHeight(34)
        self.sub_fmt_combo.addItems(["同原格式", "SRT (.srt)", "VTT (.vtt)", "ASS (.ass)"])
        mode_row.addWidget(self.sub_fmt_combo)
        mode_row.addStretch()
        sl_layout.addLayout(mode_row)

        # 字幕预览
        self.sub_preview = QTextEdit()
        self.sub_preview.setReadOnly(True)
        self.sub_preview.setPlaceholderText("导入字幕文件后，此处显示预览...")
        self.sub_preview.setMinimumHeight(120)
        sl_layout.addWidget(self.sub_preview)

        # 字幕操作按钮
        sub_action_row = QHBoxLayout()
        self.sub_translate_btn = QPushButton("  翻译字幕  ")
        self.sub_translate_btn.setObjectName("primaryBtn")
        self.sub_translate_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.sub_translate_btn.setMinimumHeight(36)
        self.sub_translate_btn.clicked.connect(self.do_subtitle_translate)
        self.sub_translate_btn.setEnabled(False)
        self.sub_export_btn = QPushButton("  导出字幕  ")
        self.sub_export_btn.setObjectName("outlineBtn")
        self.sub_export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.sub_export_btn.setMinimumHeight(36)
        self.sub_export_btn.clicked.connect(self.do_subtitle_export)
        self.sub_export_btn.setEnabled(False)
        sub_action_row.addWidget(self.sub_translate_btn, stretch=0)
        sub_action_row.addSpacing(8)
        sub_action_row.addWidget(self.sub_export_btn)
        sub_action_row.addStretch()
        sl_layout.addLayout(sub_action_row)

        self.tabs.addTab(sw, "  字幕翻译  ")

        il.addWidget(self.tabs)

        layout.addWidget(ig)

        # ── 操作栏 ──
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self.translate_btn = QPushButton("  开始翻译  ")
        self.translate_btn.setObjectName("primaryBtn")
        self.translate_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.translate_btn.setMinimumHeight(40)
        self.translate_btn.clicked.connect(self.do_translate)

        self.export_btn = QPushButton("  导出译文  ")
        self.export_btn.setObjectName("outlineBtn")
        self.export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.export_btn.setMinimumHeight(36)
        self.export_btn.setEnabled(False)
        self.export_btn.clicked.connect(self.do_export)

        self.export_term_btn = QPushButton("  导出术语  ")
        self.export_term_btn.setObjectName("outlineBtn")
        self.export_term_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.export_term_btn.setMinimumHeight(36)
        self.export_term_btn.clicked.connect(self.export_glossary)

        self.export_critique_btn = QPushButton("  导出审校  ")
        self.export_critique_btn.setObjectName("outlineBtn")
        self.export_critique_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.export_critique_btn.setMinimumHeight(36)
        self.export_critique_btn.setEnabled(False)
        self.export_critique_btn.clicked.connect(self.export_critique)

        self.glossary_btn = QPushButton("  术语库  ")
        self.glossary_btn.setObjectName("outlineBtn")
        self.glossary_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.glossary_btn.setMinimumHeight(36)
        self.glossary_btn.clicked.connect(self.open_glossary_manager)

        btn_row.addWidget(self.translate_btn, stretch=0)
        btn_row.addSpacing(12)
        btn_row.addWidget(self.export_btn)
        btn_row.addWidget(self.export_term_btn)
        btn_row.addWidget(self.export_critique_btn)
        btn_row.addWidget(self.glossary_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        # ── 进度条 ──
        progress_container = QFrame()
        progress_container.setStyleSheet("background: transparent; border: none;")
        pc_layout = QVBoxLayout(progress_container)
        pc_layout.setSpacing(6)
        pc_layout.setContentsMargins(0, 0, 0, 0)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 5)
        self.progress_bar.setValue(0)
        self.progress_bar.setFixedHeight(8)
        self.progress_bar.setTextVisible(False)
        pc_layout.addWidget(self.progress_bar)

        self.progress_label = QLabel("")
        self.progress_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.progress_label.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; background: transparent; border: none;")
        pc_layout.addWidget(self.progress_label)
        layout.addWidget(progress_container)

        # ── 结果卡片 ──
        rg = QFrame()
        rg.setObjectName("resultCard")
        rg.setStyleSheet(f"""
            QFrame#resultCard {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_LG};
            }}
        """)
        make_shadow_widget(rg, "sm")
        rl = QVBoxLayout(rg)
        rl.setSpacing(8)
        rl.setContentsMargins(20, 16, 20, 16)

        result_title = QLabel("翻译结果")
        result_title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 13, QFont.Weight.Bold))
        result_title.setStyleSheet(f"color: {C.TEXT_PRIMARY}; border: none; background: transparent;")
        rl.addWidget(result_title)

        self.result_tabs = QTabWidget()
        self.result_tabs.setStyleSheet(f"""
            QTabWidget::pane {{
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_MD};
                background: {C.BG_CARD};
                padding: 4px;
            }}
        """)

        self.analysis_output = QTextEdit()
        self.analysis_output.setReadOnly(True)
        self.result_tabs.addTab(self.analysis_output, "  分析报告  ")

        self.critique_output = QTextEdit()
        self.critique_output.setReadOnly(True)
        self.result_tabs.addTab(self.critique_output, "  审校报告  ")

        fww = QWidget()
        fwl = QVBoxLayout(fww)
        fwl.setContentsMargins(8, 8, 8, 8)
        self.result_output = QTextEdit()
        self.result_output.setReadOnly(True)
        fwl.addWidget(self.result_output)
        copy_btn = QPushButton("  复制译文  ")
        copy_btn.setObjectName("textBtn")
        copy_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        copy_btn.clicked.connect(self.copy_result)
        fwl.addWidget(copy_btn)
        self.result_tabs.addTab(fww, "  终稿译文  ")

        rl.addWidget(self.result_tabs)

        layout.addWidget(rg, stretch=1)

    # ═══════════════════════════════════════════════════════════
    # OCR 进度回调
    # ═══════════════════════════════════════════════════════════
    def _ocr_progress_handler(self, current, total, message):
        from PyQt6.QtCore import QMetaObject, Q_ARG
        QMetaObject.invokeMethod(
            self.progress_label, "setText",
            Q_ARG(str, message)
        )
        if total > 0:
            QMetaObject.invokeMethod(
                self.progress_bar, "setValue",
                Q_ARG(int, int(current * 5 / max(total, 1)))
            )
        QApplication.processEvents()

    # ═══════════════════════════════════════════════════════════
    # 文件操作
    # ═══════════════════════════════════════════════════════════
    def choose_file(self):
        try:
            path, _ = QFileDialog.getOpenFileName(
                filter="所有支持格式 (*.pdf *.docx *.doc *.xlsx *.xls *.pptx *.ppt)"
            )
            if not path:
                return

            from file_handler import read_file, set_ocr_progress_callback

            ext = path.lower().split(".")[-1]
            if ext == "pdf":
                set_ocr_progress_callback(self._ocr_progress_handler)
                self.progress_label.setText("正在读取 PDF...")
                self.progress_bar.setRange(0, 5)
                self.progress_bar.setValue(0)
                QApplication.processEvents()

            file_type, content, extra = read_file(path)

            set_ocr_progress_callback(None)
            self.progress_bar.setValue(0)
            self.progress_label.setText("")

            self.text_input.setPlainText(content)

            self.file_label.setText(f"已加载：{os.path.basename(path)}")
            self.file_path = path
            self.file_type = file_type
            self.file_extra = extra

        except Exception as e:
            QMessageBox.critical(self, "错误", f"读取失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 翻译流程
    # ═══════════════════════════════════════════════════════════
    def do_translate(self):
        tab = self.tabs.currentIndex()
        source_text = ""
        try:
            if tab == 0:
                source_text = self.text_input.toPlainText().strip()
                if not source_text:
                    QMessageBox.warning(self, "提示", "请输入需要翻译的文字。")
                    return
            elif tab == 1:
                url = self.url_input.text().strip()
                if not url:
                    QMessageBox.warning(self, "提示", "请输入URL。")
                    return
                source_text = fetch_url(url)
            elif tab == 2:
                if not self.file_path:
                    QMessageBox.warning(self, "提示", "请选择文件。")
                    return
                source_text = self.text_input.toPlainText().strip()
        except Exception as e:
            QMessageBox.critical(self, "错误", str(e))
            return

        self.last_source = source_text
        self.translate_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        self.export_term_btn.setEnabled(False)
        self.export_critique_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        self.analysis_output.clear()
        self.critique_output.clear()
        self.result_output.clear()

        self.worker = TranslateWorker(
            source_text=source_text,
            source_lang=self._get_source_lang(),
            target_lang=self._get_target_lang(),
            style=self.style_input.text(),
            audience=self.audience_input.text(),
        )
        self.worker.progress.connect(self.on_progress)
        self.worker.analysis_done.connect(self.on_analysis_done)
        self.worker.critique_done.connect(self.on_critique_done)
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_progress(self, msg, step):
        self.progress_label.setText(msg)
        self.progress_bar.setValue(step)

    def on_analysis_done(self, analysis):
        self.last_analysis = analysis
        self.analysis_output.setPlainText(analysis)
        self.export_term_btn.setEnabled(True)
        # 自动切换到分析报告tab让用户看到
        self.result_tabs.setCurrentIndex(0)

    def on_critique_done(self, critique):
        self.last_critique = critique
        self.critique_output.setPlainText(critique)
        self.export_critique_btn.setEnabled(True)

    def on_finished(self, result):
        self.last_result = result
        self.result_output.setPlainText(result)
        self.progress_label.setText("翻译完成！")
        self.progress_bar.setValue(5)
        self.translate_btn.setEnabled(True)
        self.export_btn.setEnabled(True)
        self.result_tabs.setCurrentIndex(2)

    def on_error(self, msg):
        QMessageBox.critical(self, "错误", msg)
        self.progress_label.setText("")
        self.translate_btn.setEnabled(True)

    def _get_source_lang(self):
        return self.source_lang.currentText() if hasattr(self, 'source_lang') and self.source_lang else "English"

    def _get_target_lang(self):
        return self.target_lang.currentText() if hasattr(self, 'target_lang') and self.target_lang else "Chinese"

    # ═══════════════════════════════════════════════════════════
    # 术语导出
    # ═══════════════════════════════════════════════════════════
    def export_glossary(self):
        """从分析报告+审校报告提取术语并导出"""
        from glossary import get_terms, extract_terms_from_analysis, add_terms_batch

        source_lang = self._get_source_lang()
        target_lang = self._get_target_lang()

        # 先确保术语已从分析报告提取入库
        if self.last_analysis:
            new_terms = extract_terms_from_analysis(self.last_analysis, source_lang, target_lang)
            if new_terms:
                added, _ = add_terms_batch(new_terms, source_lang, target_lang)
                if added > 0:
                    QMessageBox.information(self, "提示", f"从分析报告提取并新增了 {added} 条术语。")

        terms = get_terms(source_lang, target_lang)
        if not terms:
            QMessageBox.information(self, "提示",
                                    f"当前语言对（{source_lang} → {target_lang}）的术语库为空。\n\n"
                                    "请先完成一次翻译，术语会自动从分析报告中提取。")
            return

        # 选择导出格式
        path, _ = QFileDialog.getSaveFileName(
            self, "导出术语库", f"术语表_{source_lang}_{target_lang}",
            "Excel 文档 (*.xlsx);;Word 文档 (*.docx);;文本文件 (*.txt);;JSON 文件 (*.json)")
        if not path:
            return

        try:
            ext = path.lower().split(".")[-1]

            if ext == "xlsx":
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "术语表"

                # 标题行样式
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                header_font = Font(bold=True, size=12, color="FFFFFF")
                header_fill = PatternFill(start_color="165DFF", end_color="165DFF", fill_type="solid")
                thin_border = Border(
                    left=Side(style='thin', color='E5E6EB'),
                    right=Side(style='thin', color='E5E6EB'),
                    top=Side(style='thin', color='E5E6EB'),
                    bottom=Side(style='thin', color='E5E6EB'))

                headers = ["原文", "译文", "分类"]
                for col, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border

                # 按分类分组
                categories = {}
                for t in terms:
                    cat = t.get("category", "通用")
                    categories.setdefault(cat, []).append(t)

                row = 2
                for cat, items in categories.items():
                    # 分类标题行
                    cell = ws.cell(row=row, column=1, value=f"【{cat}】")
                    cell.font = Font(bold=True, size=11, color="165DFF")
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
                    row += 1
                    for t in items:
                        ws.cell(row=row, column=1, value=t.get("source", "")).border = thin_border
                        ws.cell(row=row, column=2, value=t.get("target", "")).border = thin_border
                        ws.cell(row=row, column=3, value=t.get("category", "")).border = thin_border
                        row += 1

                # 列宽
                ws.column_dimensions['A'].width = 30
                ws.column_dimensions['B'].width = 30
                ws.column_dimensions['C'].width = 15

                wb.save(path)

            elif ext == "docx":
                from docx import Document
                from docx.shared import Pt, RGBColor, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()

                # 标题
                title = doc.add_heading(f"术语表 ({source_lang} → {target_lang})", level=1)

                # 按分类分组
                categories = {}
                for t in terms:
                    cat = t.get("category", "通用")
                    categories.setdefault(cat, []).append(t)

                for cat, items in categories.items():
                    doc.add_heading(cat, level=2)

                    # 表格
                    table = doc.add_table(rows=1, cols=2, style='Table Grid')
                    hdr = table.rows[0].cells
                    hdr[0].text = source_lang
                    hdr[1].text = target_lang
                    for cell in hdr:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.bold = True

                    for t in items:
                        row = table.add_row()
                        row.cells[0].text = t.get("source", "")
                        row.cells[1].text = t.get("target", "")

                doc.save(path)

            elif ext == "json":
                from glossary import export_glossary_json
                content = export_glossary_json(source_lang, target_lang)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)

            else:  # txt
                from glossary import export_glossary_text
                content = export_glossary_text(source_lang, target_lang)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)

            QMessageBox.information(self, "完成", f"术语表已导出到：\n{path}\n共 {len(terms)} 条术语。")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 审校导出
    # ═══════════════════════════════════════════════════════════
    def export_critique(self):
        if not self.last_critique:
            QMessageBox.warning(self, "提示", "暂无审校报告可导出。")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "导出审校报告", "审校报告.txt",
            "文本文件 (*.txt);;Markdown (*.md)")
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.last_critique)
            QMessageBox.information(self, "完成", f"审校报告已导出到：\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 导出译文
    # ═══════════════════════════════════════════════════════════
    def do_export(self):
        if not self.last_result:
            QMessageBox.warning(self, "提示", "暂无译文可导出。")
            return
        dlg = ExportDialog(self, self.file_type)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        fmt = dlg.get_fmt()
        mode = dlg.get_mode()
        try:
            if fmt in ("docx", "doc"):
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                if mode == "bilingual":
                    doc.add_heading("双语对照", level=1)
                    src_lines = self.last_source.split("\n")
                    tgt_lines = self.last_result.split("\n")
                    for s, t in zip(src_lines, tgt_lines):
                        p = doc.add_paragraph()
                        p.add_run(s).font.size = Pt(10)
                        p = doc.add_paragraph()
                        run = p.add_run(t)
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x16, 0x5D, 0xFF)
                elif mode == "translation":
                    doc.add_paragraph(self.last_result)
                else:
                    doc.add_paragraph(self.last_result)
                path, _ = QFileDialog.getSaveFileName(
                    self, "保存", "译文.docx", "Word 文档 (*.docx)")
                if path:
                    doc.save(path)
            elif fmt in ("xlsx", "xls"):
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "译文"
                ws.append(["原文", "译文"])
                src_lines = self.last_source.split("\n")
                tgt_lines = self.last_result.split("\n")
                for s, t in zip(src_lines, tgt_lines):
                    ws.append([s, t])
                path, _ = QFileDialog.getSaveFileName(
                    self, "保存", "译文.xlsx", "Excel 文档 (*.xlsx)")
                if path:
                    wb.save(path)
            else:
                path, _ = QFileDialog.getSaveFileName(
                    self, "保存", "译文.txt", "文本文件 (*.txt)")
                if path:
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(self.last_result)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 复制结果
    # ═══════════════════════════════════════════════════════════
    def copy_result(self):
        if self.last_result:
            QApplication.clipboard().setText(self.last_result)
            QMessageBox.information(self, "完成", "译文已复制到剪贴板。")

    # ═══════════════════════════════════════════════════════════
    # 字幕导入（含自动清洗）
    # ═══════════════════════════════════════════════════════════
    def choose_subtitle_file(self):
        try:
            path, _ = QFileDialog.getOpenFileName(
                filter="字幕文件 (*.srt *.vtt *.ass *.ssa)")
            if not path:
                return

            from subtitle_handler import (
                parse_subtitle_file,
                clean_subtitle_file,
            )

            parsed = parse_subtitle_file(path)
            if not parsed or not parsed.entries:
                QMessageBox.warning(self, "提示", "无法解析该字幕文件。")
                return

            # ── 自动清洗字幕 ──
            parsed = clean_subtitle_file(parsed)

            self.subtitle_obj = parsed
            self.subtitle_path = path
            self.sub_file_label.setText(f"已导入（已清洗）：{os.path.basename(path)} [{len(parsed.entries)} 条]")
            self.sub_translate_btn.setEnabled(True)
            self.sub_export_btn.setEnabled(False)

            # 预览前 20 条
            preview_lines = []
            for e in parsed.entries[:20]:
                preview_lines.append(f"[{e.index}] {e.start_time} --> {e.end_time}")
                preview_lines.append(f"    {e.text}")
                preview_lines.append("")
            if len(parsed.entries) > 20:
                preview_lines.append(f"... 共 {len(parsed.entries)} 条，仅显示前 20 条")
            self.sub_preview.setPlainText("\n".join(preview_lines))

            # 自动设置输出格式
            ext = path.lower().split(".")[-1]
            fmt_map = {"srt": "SRT (.srt)", "vtt": "VTT (.vtt)", "ass": "ASS (.ass)", "ssa": "ASS (.ass)"}
            target_fmt = fmt_map.get(ext, "同原格式")
            idx = self.sub_fmt_combo.findText(target_fmt)
            if idx >= 0:
                self.sub_fmt_combo.setCurrentIndex(idx)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"导入字幕失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 字幕翻译
    # ═══════════════════════════════════════════════════════════
    def do_subtitle_translate(self):
        if not self.subtitle_obj:
            QMessageBox.warning(self, "提示", "请先导入字幕文件。")
            return
        self.sub_translate_btn.setEnabled(False)
        self.sub_export_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_label.setText("")

        self.subtitle_worker = SubtitleWorker(
            self.subtitle_obj,
            self._get_source_lang(),
            self._get_target_lang(),
            style=self.style_input.text() if hasattr(self, 'style_input') else "",
            audience=self.audience_input.text() if hasattr(self, 'audience_input') else "",
        )
        self.subtitle_worker.progress.connect(self._on_subtitle_progress)
        self.subtitle_worker.analysis_done.connect(self._on_subtitle_analysis_done)
        self.subtitle_worker.critique_done.connect(self._on_subtitle_critique_done)
        self.subtitle_worker.finished.connect(self._on_subtitle_finished)
        self.subtitle_worker.error.connect(self._on_subtitle_error)
        self.subtitle_worker.start()

    def _on_subtitle_progress(self, msg, step):
        self.progress_label.setText(msg)
        self.progress_bar.setValue(step)

    def _on_subtitle_error(self, msg):
        QMessageBox.critical(self, "字幕翻译错误", msg)
        self.progress_label.setText("")
        self.sub_translate_btn.setEnabled(True)
        self.sub_export_btn.setEnabled(True)

    def _on_subtitle_analysis_done(self, analysis):
        """第一步完成：实时显示 AI 分析报告到分析 tab"""
        self.last_analysis = analysis
        self.analysis_output.setPlainText(analysis)
        self.export_term_btn.setEnabled(True)
        self.result_tabs.setCurrentIndex(0)

    def _on_subtitle_critique_done(self, critique):
        """第四步完成：实时显示 AI 审校报告到审校 tab"""
        self.last_critique = critique
        self.critique_output.setPlainText(critique)
        self.export_critique_btn.setEnabled(True)

    def _on_subtitle_finished(self, subtitle, translated):
        """第五步完成：更新字幕对象、填充结果 tab"""
        self.subtitle_obj = subtitle
        self.sub_translate_btn.setEnabled(True)
        self.sub_export_btn.setEnabled(True)
        self.progress_label.setText("字幕翻译完成！")
        self.progress_bar.setValue(5)

        source_lang = self._get_source_lang()
        target_lang = self._get_target_lang()
        total = len(subtitle.entries)
        translated_count = sum(1 for e in subtitle.entries if e.translated.strip())

        # 填充结果 tab（终稿译文）
        result_lines = []
        for e in subtitle.entries:
            if e.translated:
                result_lines.append(e.translated)
        self.result_output.setPlainText("\n".join(result_lines))
        self.export_btn.setEnabled(True)
        self.result_tabs.setCurrentIndex(2)

        # 更新字幕预览
        preview_lines = []
        for e in subtitle.entries[:20]:
            preview_lines.append(f"[{e.index}] {e.start_time} --> {e.end_time}")
            preview_lines.append(f"    {e.text}")
            if e.translated:
                preview_lines.append(f"    → {e.translated}")
            preview_lines.append("")
        if len(subtitle.entries) > 20:
            preview_lines.append(f"... 共 {len(subtitle.entries)} 条")
        self.sub_preview.setPlainText("\n".join(preview_lines))
    # ═══════════════════════════════════════════════════════════
    # 字幕导出
    # ═══════════════════════════════════════════════════════════
    def do_subtitle_export(self):
        if not self.subtitle_obj:
            QMessageBox.warning(self, "提示", "没有可导出的字幕。")
            return
        try:
            from subtitle_handler import (
                export_subtitle_file,
                optimize_subtitle_for_video,
            )

            mode = self.sub_mode_combo.currentIndex()
            fmt_text = self.sub_fmt_combo.currentText()

            if "SRT" in fmt_text:
                out_fmt = "srt"
            elif "VTT" in fmt_text:
                out_fmt = "vtt"
            elif "ASS" in fmt_text:
                out_fmt = "ass"
            else:
                ext = self.subtitle_path.lower().split(".")[-1] if self.subtitle_path else "srt"
                out_fmt = ext if ext in ("srt", "vtt", "ass", "ssa") else "srt"

            base = os.path.splitext(os.path.basename(self.subtitle_path))[0]
            mode_names = ["双语", "纯译文", "Clean纯文本", "Clean双语"]
            default_name = f"{base}_{mode_names[mode]}.{out_fmt}"

            path, _ = QFileDialog.getSaveFileName(
                self, "导出字幕", default_name,
                f"{out_fmt.upper()} 文件 (*.{out_fmt})")
            if not path:
                return

            # 优化适配视频
            try:
                optimized = optimize_subtitle_for_video(self.subtitle_obj)
            except Exception:
                optimized = self.subtitle_obj

            export_subtitle_file(
                optimized, path, output_mode=mode, output_format=out_fmt)

            QMessageBox.information(self, "完成",
                                    f"字幕已导出到：\n{path}\n已自动优化：智能换行 + 时长适配")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")

    # ═══════════════════════════════════════════════════════════
    # 术语库管理器
    # ═══════════════════════════════════════════════════════════
    def open_glossary_manager(self):
        from glossary import (
            get_terms, add_term, edit_term, delete_term,
            clear_glossary, import_glossary_from_text, get_all_lang_pairs,
        )
        C = ArcoColors

        dlg = QDialog(self)
        dlg.setWindowTitle("术语库管理")
        dlg.setMinimumSize(780, 600)
        dlg.resize(820, 650)

        # ── 全局样式 ──
        dlg.setStyleSheet(f"""
            QDialog {{
                background-color: {C.BG_PAGE};
                color: {C.TEXT_PRIMARY};
            }}
            QLabel {{
                color: {C.TEXT_PRIMARY};
                background: transparent;
            }}
            QLineEdit {{
                color: {C.TEXT_PRIMARY};
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER};
                border-radius: {C.RADIUS_SM};
                padding: 0 8px;
                font-size: 13px;
            }}
            QLineEdit:focus {{
                border-color: {C.PRIMARY};
            }}
            QPushButton {{
                color: {C.TEXT_PRIMARY};
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER};
                border-radius: {C.RADIUS_SM};
                padding: 6px 16px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                border-color: {C.PRIMARY};
                color: {C.PRIMARY};
            }}
            QPushButton:pressed {{
                background: {C.PRIMARY_LIGHT};
            }}
            QTableWidget {{
                background-color: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_MD};
                gridline-color: {C.BORDER_LIGHT};
                font-size: 13px;
                selection-background-color: {C.PRIMARY_LIGHT};
                selection-color: {C.TEXT_PRIMARY};
                alternate-background-color: {C.FILL1};
            }}
            QTableWidget::item {{
                padding: 6px 8px;
                border-bottom: 1px solid {C.BORDER_LIGHT};
            }}
            QTableWidget::item:hover {{
                background-color: {C.PRIMARY_LIGHT2};
            }}
            QHeaderView::section {{
                background-color: {C.FILL2};
                color: {C.TEXT_REGULAR};
                border: none;
                border-bottom: 2px solid {C.BORDER};
                border-right: 1px solid {C.BORDER_LIGHT};
                padding: 8px 6px;
                font-size: 13px;
                font-weight: 600;
            }}
            QComboBox {{
                color: {C.TEXT_PRIMARY};
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER};
                border-radius: {C.RADIUS_SM};
                padding: 4px 8px;
                min-height: 30px;
                font-size: 13px;
            }}
            QComboBox:focus {{
                border-color: {C.PRIMARY};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 24px;
            }}
            QComboBox QAbstractItemView {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER};
                selection-background-color: {C.PRIMARY_LIGHT};
            }}
        """)

        main_layout = QVBoxLayout(dlg)
        main_layout.setSpacing(12)
        main_layout.setContentsMargins(20, 20, 20, 20)

        # ── 标题行 ──
        title_row = QHBoxLayout()
        title = QLabel("术语库管理")
        title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 16, QFont.Weight.Bold))
        title.setStyleSheet(f"color: {C.TEXT_PRIMARY}; background: transparent;")
        title_row.addWidget(title)
        title_row.addStretch()

        # 语言对下拉框
        source_lang = self._get_source_lang()
        target_lang = self._get_target_lang()
        pair_lbl = QLabel("语言对：")
        pair_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; background: transparent; font-size: 13px;")
        title_row.addWidget(pair_lbl)

        lang_pairs = get_all_lang_pairs()
        pair_combo = QComboBox()
        pair_combo.setMinimumWidth(200)
        # 当前语言对放在第一个
        current_key = f"{source_lang.lower()[:2]}→{target_lang.lower()[:2]}"
        pair_items = []
        # 找到当前语言对的全名 key
        current_display = f"{source_lang} → {target_lang}"
        pair_items.append((current_display, source_lang, target_lang))
        for lp in lang_pairs:
            key = lp.get("key", "")
            sl = lp.get("source_lang", "")
            tl = lp.get("target_lang", "")
            if f"{sl.lower()[:2]}→{tl.lower()[:2]}" != current_key:
                pair_items.append((f"{sl} → {tl}  ({lp.get('count', 0)} 条)", sl, tl))
        pair_combo.addItems([p[0] for p in pair_items])
        title_row.addWidget(pair_combo)

        main_layout.addLayout(title_row)

        # 术语计数
        count_lbl = QLabel("")
        count_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; background: transparent; font-size: 12px;")
        main_layout.addWidget(count_lbl)

        # ── 搜索栏 ──
        search_frame = QFrame()
        search_frame.setStyleSheet(f"""
            QFrame {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_MD};
            }}
        """)
        search_layout = QHBoxLayout(search_frame)
        search_layout.setContentsMargins(10, 6, 10, 6)
        search_layout.setSpacing(8)
        search_icon = QLabel("🔍")
        search_icon.setStyleSheet("background: transparent; font-size: 14px;")
        search_layout.addWidget(search_icon)
        search_le = QLineEdit()
        search_le.setPlaceholderText("搜索术语（原文或译文）...")
        search_le.setFrame(False)  # 去掉默认边框
        search_le.setFixedHeight(30)
        search_layout.addWidget(search_le)
        main_layout.addWidget(search_frame)

        # ── 添加术语区 ──
        add_frame = QFrame()
        add_frame.setStyleSheet(f"""
            QFrame {{
                background: {C.BG_CARD};
                border: 1px solid {C.BORDER_LIGHT};
                border-radius: {C.RADIUS_MD};
            }}
        """)
        add_layout = QHBoxLayout(add_frame)
        add_layout.setContentsMargins(12, 10, 12, 10)
        add_layout.setSpacing(8)

        src_label = QLabel("原文")
        src_label.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; background: transparent;")
        src_le = QLineEdit()
        src_le.setPlaceholderText("输入原文术语")
        src_le.setFixedHeight(34)
        add_layout.addWidget(src_label)
        add_layout.addWidget(src_le, stretch=3)

        arrow_lbl = QLabel("→")
        arrow_lbl.setStyleSheet(f"color: {C.PRIMARY}; font-size: 16px; font-weight: bold; background: transparent;")
        arrow_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        add_layout.addWidget(arrow_lbl)

        tgt_label = QLabel("译文")
        tgt_label.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 12px; background: transparent;")
        tgt_le = QLineEdit()
        tgt_le.setPlaceholderText("输入译文")
        tgt_le.setFixedHeight(34)
        add_layout.addWidget(tgt_label)
        add_layout.addWidget(tgt_le, stretch=3)

        cat_le = QLineEdit()
        cat_le.setPlaceholderText("分类")
        cat_le.setFixedWidth(80)
        cat_le.setFixedHeight(34)
        add_layout.addWidget(cat_le)

        add_btn = QPushButton("＋ 添加")
        add_btn.setFixedHeight(34)
        add_btn.setMinimumWidth(70)
        add_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        add_btn.setStyleSheet(f"""
            QPushButton {{
                background: {C.PRIMARY};
                color: white;
                border: none;
                border-radius: {C.RADIUS_SM};
                font-weight: 600;
            }}
            QPushButton:hover {{
                background: {C.PRIMARY_HOVER};
                color: white;
                border: none;
            }}
            QPushButton:pressed {{
                background: {C.PRIMARY_ACTIVE};
                color: white;
                border: none;
            }}
        """)
        add_layout.addWidget(add_btn)
        main_layout.addWidget(add_frame)

        # ── 术语表格 ──
        table = QTableWidget(0, 4)
        table.setHorizontalHeaderLabels(["原文", "译文", "分类", "添加时间"])
        table.horizontalHeader().setStretchLastSection(True)
        table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Fixed)
        table.setColumnWidth(2, 80)
        table.setColumnWidth(3, 110)
        table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        table.setAlternatingRowColors(True)
        table.verticalHeader().setVisible(False)
        table.setShowGrid(False)
        main_layout.addWidget(table)

        # ── 状态栏 + 提示 ──
        hint_lbl = QLabel("💡 双击表格行可编辑术语")
        hint_lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; font-size: 11px; background: transparent;")
        main_layout.addWidget(hint_lbl)

        # ── 底部按钮 ──
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_configs = [
            ("📥 导入术语", do_import := lambda: None),  # placeholder
            ("✏️ 编辑选中", do_edit := lambda: None),
            ("🗑️ 删除选中", do_delete := lambda: None),
            ("清空全部", do_clear := lambda: None),
        ]
        action_btns = {}
        for text, _ in btn_configs:
            b = QPushButton(text)
            b.setFixedHeight(34)
            b.setCursor(Qt.CursorShape.PointingHandCursor)
            action_btns[text] = b
            btn_row.addWidget(b)
        close_btn = QPushButton("关闭")
        close_btn.setFixedHeight(34)
        close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        close_btn.clicked.connect(dlg.accept)
        btn_row.addWidget(close_btn)
        main_layout.addLayout(btn_row)

        # ── 数据与逻辑 ──
        def get_current_langs():
            """获取当前选择的语言对"""
            idx = pair_combo.currentIndex()
            if 0 <= idx < len(pair_items):
                return pair_items[idx][1], pair_items[idx][2]
            return source_lang, target_lang

        def refresh_table(search_text=""):
            sl, tl = get_current_langs()
            terms = get_terms(sl, tl)
            # 搜索过滤
            if search_text.strip():
                kw = search_text.strip().lower()
                terms = [t for t in terms
                         if kw in t.get("source", "").lower()
                         or kw in t.get("target", "").lower()
                         or kw in t.get("category", "").lower()]
            table.setRowCount(len(terms))
            for i, t in enumerate(terms):
                table.setItem(i, 0, QTableWidgetItem(t.get("source", "")))
                table.setItem(i, 1, QTableWidgetItem(t.get("target", "")))
                table.setItem(i, 2, QTableWidgetItem(t.get("category", "")))
                table.setItem(i, 3, QTableWidgetItem(t.get("added_at", "")))
            count_lbl.setText(f"共 {len(terms)} 条术语")

        def do_add_term():
            sl, tl = get_current_langs()
            s = src_le.text().strip()
            t = tgt_le.text().strip()
            c = cat_le.text().strip() or "通用"
            if not s or not t:
                QMessageBox.warning(dlg, "提示", "请输入原文和译文。")
                return
            result = add_term(s, t, sl, tl, category=c)
            if result:
                src_le.clear(); tgt_le.clear(); cat_le.clear()
                src_le.setFocus()
                refresh_table(search_le.text())
            else:
                QMessageBox.information(dlg, "提示",
                    f"术语 \"{s}\" 已存在。\n如需修改，请双击该术语或选中后点击编辑按钮。")

        def do_edit_term():
            sl, tl = get_current_langs()
            rows = set(i.row() for i in table.selectedItems())
            if not rows:
                QMessageBox.warning(dlg, "提示", "请先选择要编辑的术语。")
                return
            terms = get_terms(sl, tl)
            kw = search_le.text().strip().lower()
            if kw:
                terms = [t for t in terms
                         if kw in t.get("source", "").lower()
                         or kw in t.get("target", "").lower()]
            row = min(rows)
            if row >= len(terms):
                return
            t = terms[row]
            # 弹出编辑对话框
            edit_dlg = QDialog(dlg)
            edit_dlg.setWindowTitle("编辑术语")
            edit_dlg.setFixedSize(420, 200)
            edit_dlg.setStyleSheet(dlg.styleSheet())
            edit_layout = QVBoxLayout(edit_dlg)
            edit_layout.setSpacing(10)
            edit_layout.setContentsMargins(20, 20, 20, 20)

            edit_title = QLabel("编辑术语")
            edit_title.setFont(QFont(C.FONT_FAMILY.split(",")[0].strip().strip("'"), 14, QFont.Weight.Bold))
            edit_layout.addWidget(edit_title)

            form = QGridLayout()
            form.setSpacing(8)
            for col, (label_text, default_val) in enumerate([
                ("原文：", t.get("source", "")),
                ("译文：", t.get("target", "")),
                ("分类：", t.get("category", "")),
            ]):
                lbl = QLabel(label_text)
                lbl.setStyleSheet(f"color: {C.TEXT_SECONDARY}; background: transparent;")
                le = QLineEdit(default_val)
                le.setFixedHeight(32)
                form.addWidget(lbl, col, 0)
                form.addWidget(le, col, 1)
            edit_layout.addLayout(form)

            btn_row2 = QHBoxLayout()
            btn_row2.addStretch()
            save_btn = QPushButton("保存")
            save_btn.setFixedHeight(32)
            save_btn.setMinimumWidth(80)
            save_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            save_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {C.PRIMARY}; color: white; border: none;
                    border-radius: {C.RADIUS_SM}; font-weight: 600;
                }}
                QPushButton:hover {{ background: {C.PRIMARY_HOVER}; color: white; border: none; }}
            """)
            cancel_btn = QPushButton("取消")
            cancel_btn.setFixedHeight(32)
            cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn_row2.addWidget(save_btn)
            btn_row2.addWidget(cancel_btn)
            edit_layout.addLayout(btn_row2)

            inputs = [form.itemAt(r, 1).widget() for r in range(3)]

            def do_save():
                new_src = inputs[0].text().strip()
                new_tgt = inputs[1].text().strip()
                new_cat = inputs[2].text().strip()
                if not new_src or not new_tgt:
                    QMessageBox.warning(edit_dlg, "提示", "原文和译文不能为空。")
                    return
                ok = edit_term(sl, tl, t.get("source", ""),
                               new_target=new_tgt,
                               new_source=new_src,
                               new_category=new_cat)
                if ok:
                    edit_dlg.accept()
                    refresh_table(search_le.text())
                else:
                    QMessageBox.warning(edit_dlg, "失败", "编辑失败，术语可能已被删除。")

            save_btn.clicked.connect(do_save)
            cancel_btn.clicked.connect(edit_dlg.reject)
            edit_dlg.exec()

        def do_delete_terms():
            sl, tl = get_current_langs()
            rows = set(i.row() for i in table.selectedItems())
            if not rows:
                QMessageBox.warning(dlg, "提示", "请先选择要删除的术语。")
                return
            terms = get_terms(sl, tl)
            kw = search_le.text().strip().lower()
            if kw:
                terms = [t for t in terms
                         if kw in t.get("source", "").lower()
                         or kw in t.get("target", "").lower()]
            deleted = 0
            for row in sorted(rows, reverse=True):
                if row < len(terms):
                    t = terms[row]
                    if delete_term(sl, tl, t.get("source", "")):
                        deleted += 1
            refresh_table(search_le.text())

        def do_clear_all():
            sl, tl = get_current_langs()
            ret = QMessageBox.question(dlg, "确认",
                                       f"确定要清空 {sl} → {tl} 的所有术语吗？\n此操作不可撤销。",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if ret == QMessageBox.StandardButton.Yes:
                clear_glossary(sl, tl)
                refresh_table(search_le.text())

        def do_import_terms():
            sl, tl = get_current_langs()
            path, _ = QFileDialog.getOpenFileName(dlg, "导入术语", "", "文本文件 (*.txt)")
            if path:
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        content = f.read()
                    added, total = import_glossary_from_text(content, sl, tl)
                    refresh_table(search_le.text())
                    QMessageBox.information(dlg, "完成",
                                            f"导入完成：共 {total} 条，新增 {added} 条，跳过 {total - added} 条。")
                except Exception as e:
                    QMessageBox.critical(dlg, "错误", f"导入失败：{str(e)}")

        # 双击编辑
        def on_double_click(row, _col):
            table.selectRow(row)
            do_edit_term()

        # 连接信号
        add_btn.clicked.connect(do_add_term)
        action_btns["✏️ 编辑选中"].clicked.connect(do_edit_term)
        action_btns["🗑️ 删除选中"].clicked.connect(do_delete_terms)
        action_btns["清空全部"].clicked.connect(do_clear_all)
        action_btns["📥 导入术语"].clicked.connect(do_import_terms)
        table.cellDoubleClicked.connect(on_double_click)
        pair_combo.currentIndexChanged.connect(lambda: refresh_table(search_le.text()))
        search_le.textChanged.connect(refresh_table)

        # Enter 键添加
        def on_enter():
            do_add_term()
        tgt_le.returnPressed.connect(on_enter)
        cat_le.returnPressed.connect(on_enter)
        src_le.returnPressed.connect(lambda: tgt_le.setFocus())

        # 初始加载
        refresh_table()
        dlg.exec()


# ═══════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    win = MainWindow()
    win.show()
    sys.exit(app.exec())